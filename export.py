import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_word(entries: list[tuple[str, dict]]) -> io.BytesIO:
    doc = Document()

    # remove default margins
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Pt(40)
    section.left_margin = section.right_margin = Pt(50)

    for i, (date_key, entry) in enumerate(entries):
        display = ".".join(date_key.split(".")[:2])
        foods = ", ".join(entry["foods"]) if entry["foods"] else "—"

        # date
        p = doc.add_paragraph()
        run = p.add_run(display)
        run.bold = True
        run.font.size = Pt(14)

        # kcal
        p = doc.add_paragraph()
        run = p.add_run(f"{entry['kcal']} ккал")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0xE0, 0x50, 0x50)

        # КБЖУ
        p = doc.add_paragraph()
        run = p.add_run(f"Б: {entry['protein']}г  Ж: {entry['fat']}г  В: {entry['carbs']}г")
        run.font.size = Pt(11)

        # foods
        p = doc.add_paragraph()
        run = p.add_run(foods)
        run.font.size = Pt(11)
        run.italic = True

        # separator between days
        if i < len(entries) - 1:
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
